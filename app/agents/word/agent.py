import os
import json
import time
import docx
from flask import Blueprint, request, jsonify, current_app
from werkzeug.utils import secure_filename
from langchain_text_splitters import RecursiveCharacterTextSplitter

word_agent_bp = Blueprint('word_agent', __name__, url_prefix='/agents/word')

@word_agent_bp.route('/scan', methods=['POST'])
def scan_word():
    """
    Scan a Word document and extract its content.
    
    This endpoint accepts a file path relative to the upload folder
    and processes the Word document to extract text, metadata, and structure.
    """
    file_path = request.json.get('file_path')
    if not file_path:
        return jsonify({'error': 'No file path provided'}), 400
        
    # Make sure the file exists and is a Word document
    full_path = os.path.join(current_app.config['UPLOAD_FOLDER'], file_path)
    if not os.path.exists(full_path):
        return jsonify({'error': 'File not found'}), 404
        
    if not file_path.lower().endswith(('.docx', '.doc')):
        return jsonify({'error': 'File is not a Word document'}), 400
    
    try:
        print(f"Starting Word document processing: {file_path}")
        start_time = time.time()
        
        # Process the document
        document_data = process_word_document(full_path)
        
        # Add processing metadata
        processing_time = time.time() - start_time
        document_data["processing_info"] = {
            "processing_time_seconds": round(processing_time, 2),
            "processed_at": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        # Convert to JSON string
        json_content = json.dumps(document_data, ensure_ascii=False, indent=2)
        
        # Create a detailed summary
        if "error" in document_data:
            summary = document_data["error"]
        else:
            # Create a text-based summary for display
            summary = f"Word Document: {os.path.basename(file_path)}\n"
            summary += f"Paragraphs: {document_data['document_info']['paragraph_count']}\n"
            
            if document_data['document_info']['table_count'] > 0:
                summary += f"Tables: {document_data['document_info']['table_count']}\n"
                
            summary += f"Size: {document_data['document_info']['file_size_kb']:.2f} KB\n"
            summary += f"Processing Time: {processing_time:.2f} seconds\n\n"
            
            # Add core properties if available
            if document_data["metadata"]:
                summary += "Document Properties:\n"
                for key, value in document_data["metadata"].items():
                    if value and value != "Not available":
                        summary += f"- {key.replace('_', ' ').title()}: {value}\n"
                summary += "\n"
            
            # Add text chunks count and extraction details
            if "chunks" in document_data:
                summary += f"Text extracted and split into {len(document_data['chunks'])} chunks for processing.\n"
                
                # Add word count if available
                total_words = sum(len(text.split()) for text in document_data.get("chunks", []))
                if total_words > 0:
                    summary += f"Total Words: Approximately {total_words}\n"
        
        # Format the summary with HTML breaks
        formatted_summary = summary.replace('\n', '<br>')
        
        # Log completion
        print(f"Completed Word document processing: {file_path} in {time.time() - start_time:.2f} seconds")
        
        return jsonify({
            'filename': os.path.basename(file_path),
            'summary': formatted_summary,
            'json_data': document_data
        })
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error processing Word document: {str(e)}\n{error_traceback}")
        
        # Return a more detailed error response
        return jsonify({
            'error': f"Failed to process Word document: {str(e)}",
            'file_path': file_path,
            'error_type': type(e).__name__
        }), 500

def process_word_document(file_path):
    """
    Process a Word document and return structured data with text chunks.
    
    This function:
    1. Extracts metadata from the document
    2. Extracts all paragraphs and their styles
    3. Extracts tables and their contents
    4. Chunks the text for further processing
    """
    try:
        # Open the document
        print(f"Opening Word document: {file_path}")
        doc = docx.Document(file_path)
        
        # Initialize document structure
        document_data = {
            "metadata": {},
            "document_info": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "file_size_kb": os.path.getsize(file_path) / 1024,
                "file_name": os.path.basename(file_path)
            },
            "content": {
                "paragraphs": [],
                "tables": [],
                "sections": []
            }
        }
        
        # Extract core properties if available
        try:
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                document_data["metadata"] = {
                    "title": props.title or "Not available",
                    "author": props.author or "Not available",
                    "subject": props.subject or "Not available",
                    "keywords": props.keywords or "Not available",
                    "category": props.category or "Not available",
                    "comments": props.comments or "Not available",
                    "created": str(props.created) if props.created else "Not available",
                    "modified": str(props.modified) if props.modified else "Not available",
                    "last_modified_by": props.last_modified_by or "Not available"
                }
        except Exception as e:
            document_data["metadata"] = {"error": f"Failed to extract metadata: {str(e)}"}
        
        # Extract sections and add to structured content
        for i, section in enumerate(doc.sections):
            section_data = {
                "index": i,
                "width": section.page_width.inches if hasattr(section, 'page_width') else None,
                "height": section.page_height.inches if hasattr(section, 'page_height') else None,
                "orientation": "portrait" if section.page_width < section.page_height else "landscape" 
                               if hasattr(section, 'page_width') and hasattr(section, 'page_height') else "unknown"
            }
            document_data["content"]["sections"].append(section_data)
        
        # Analyze styles
        styles_used = {}
        headings_count = 0
        
        # Extract paragraphs
        all_text = ""
        current_heading = None
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"
            
            # Track styles used
            if style_name not in styles_used:
                styles_used[style_name] = 0
            styles_used[style_name] += 1
            
            # Track headings
            is_heading = style_name.startswith('Heading')
            if is_heading:
                headings_count += 1
                current_heading = text
            
            if text:  # Only add non-empty paragraphs
                paragraph_data = {
                    "index": i,
                    "text": text,
                    "style": style_name,
                    "is_heading": is_heading
                }
                
                # If this paragraph is under a heading, record that
                if current_heading and not is_heading:
                    paragraph_data["under_heading"] = current_heading
                
                document_data["content"]["paragraphs"].append(paragraph_data)
                
                # Add to full text with proper formatting
                if is_heading:
                    # Add extra formatting for headings
                    all_text += f"# {text}\n\n"
                else:
                    all_text += text + "\n\n"
        
        # Add document structure statistics
        document_data["document_info"]["styles_used"] = styles_used
        document_data["document_info"]["headings_count"] = headings_count
        
        # Extract tables with enhanced data
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": [],
                "row_count": len(table.rows),
                "column_count": len(table.columns) if table.rows else 0
            }
            
            # Add all table cells as rows
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    all_text += cell_text + " "
                table_data["rows"].append(row_data)
                all_text += "\n"
            
            # Add table data to document
            document_data["content"]["tables"].append(table_data)
            all_text += "\n"
        
        # Use LangChain's RecursiveCharacterTextSplitter to chunk the text with improved settings
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1200,  # Larger chunks for more context
            chunk_overlap=250,  # More overlap for better continuity
            length_function=len,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
        # Split the text into chunks
        if all_text.strip():
            chunks = text_splitter.split_text(all_text)
            document_data["chunks"] = chunks
            
            # Add word and character counts
            total_words = sum(len(chunk.split()) for chunk in chunks)
            total_chars = sum(len(chunk) for chunk in chunks)
            document_data["document_info"]["estimated_word_count"] = total_words
            document_data["document_info"]["total_characters"] = total_chars
        else:
            document_data["chunks"] = ["No text content found in document"]
            document_data["document_info"]["estimated_word_count"] = 0
            document_data["document_info"]["total_characters"] = 0
        
        return document_data
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error processing Word document: {str(e)}\n{error_traceback}")
        return {"error": f"Error processing Word document: {str(e)}"}